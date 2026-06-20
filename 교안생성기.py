#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""소그룹교안 생성기 v1.4"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os, json, re, threading, subprocess, zipfile, sys
from pathlib import Path
from datetime import datetime
from io import BytesIO

# Windows에서 subprocess 호출 시 검은 터미널 창 숨기기
_NO_WINDOW = subprocess.CREATE_NO_WINDOW if sys.platform == 'win32' else 0

# ──────────────────────────── 텍스트 추출 ────────────────────────────────────

def extract_text(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == '.pdf':
        # 1차: pdfplumber
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        parts.append(t)
            text = '\n'.join(parts)
            if text.strip():
                return text
        except ImportError:
            raise ImportError("pdfplumber 패키지가 없습니다.\npip install pdfplumber")
        except Exception:
            pass  # 손상 PDF → 2차 시도

        # 2차: pypdf (비정상 구조 PDF 대응)
        try:
            from pypdf import PdfReader
            reader = PdfReader(path, strict=False)
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            text = '\n'.join(parts)
            if text.strip():
                return text
        except ImportError:
            pass
        except Exception:
            pass

        raise ValueError(
            "PDF 파일을 읽을 수 없습니다.\n"
            "파일이 손상되었거나 암호화된 PDF일 수 있습니다.\n\n"
            "해결 방법:\n"
            "• Chrome에서 파일 열기 → 인쇄 → 'PDF로 저장'\n"
            "• Adobe Acrobat에서 '다른 이름으로 저장' 후 재시도"
        )
    elif ext == '.docx':
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx 패키지가 없습니다.\n실행.bat 로 재실행하거나: pip install python-docx")
        return '\n'.join(p.text for p in Document(path).paragraphs)
    elif ext == '.hwpx':
        import xml.etree.ElementTree as ET
        parts = []
        with zipfile.ZipFile(path, 'r') as z:
            names = sorted(n for n in z.namelist() if re.search(r'[Ss]ection\d+\.xml', n))
            if not names:
                names = sorted(n for n in z.namelist()
                               if n.startswith('Contents/') and n.endswith('.xml'))
            for name in names:
                with z.open(name) as f:
                    try:
                        root = ET.fromstring(f.read())
                        for el in root.iter():
                            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
                            if tag.lower() == 't' and el.text:
                                parts.append(el.text)
                    except ET.ParseError:
                        pass
        return '\n'.join(parts)
    elif ext == '.hwp':
        raise ValueError(
            "구버전 HWP 형식은 직접 지원하지 않습니다.\n"
            "한컴오피스에서 [파일 → 다른 이름으로 저장 → HWPX]로 저장 후 다시 시도하세요."
        )
    else:
        for enc in ('utf-8-sig', 'utf-8', 'cp949', 'euc-kr'):
            try:
                return Path(path).read_text(encoding=enc)
            except (UnicodeDecodeError, LookupError):
                continue
        raise ValueError("파일 인코딩을 인식할 수 없습니다.")


# ──────────────────────────── 교안 포맷 ──────────────────────────────────────

SYSTEM_PROMPT = """\
당신은 설교문을 분석하여 마을예배/소그룹 교안을 작성하는 전문가입니다.

【성경 약어 — 반드시 약어 사용】
창,출,레,민,신,수,삿,룻,삼상,삼하,왕상,왕하,대상,대하,스,느,욥,시,잠,전,아,사,렘,겔,단,호,욜,암,미,합,말
마,막,눅,요,행,롬,고전,고후,갈,엡,빌,골,약,벧전,벧후,요일,계

설교문을 분석하여 아래 JSON만 출력하세요 (코드블록·설명 없이, 순수 JSON만):
{
  "title": "설교 제목 (짧게, 주제 단어)",
  "scripture": "성경 약어 본문 (예: 삿 16:15~22, 대상 21:1~8)",
  "hymn_open": "개회 찬송 번호 (설교에 언급된 것 우선, 없으면 주제에 맞는 번호)",
  "hymn_close": "폐회 찬송 번호 (개회와 다른 번호)",
  "s1": [
    "아이스브레이크 질문 1 — 설교 주제와 자연스럽게 연결, 일상·계절·감각 소재, 처음 보는 사람도 부담 없이 답할 수준",
    "아이스브레이크 질문 2"
  ],
  "life_sum1": "설교 전반부 요약. 6~9문장. 설교자 어조·표현 그대로. 본문 사건·배경 + 신학적 핵심 의미 + 히브리어·원어 디테일이나 대조 인물·핵심 책망 한두 문장 포함. 핵심 성구 본문에 녹여 인용.",
  "life_q1": "나눔 질문 1 — '최근 나의 삶에서...' 형태 + 그 경험 속 마음의 동기·내면을 묻는 보조 문장 한 줄 추가.",
  "life_sum2": "설교 후반부 요약. 6~9문장. 설교자 어조 그대로. 문제 본질 진단 + 해답·소망의 길 + 신학적 핵심 이미지(어둠/빛, 결핍/충만, 첫째 아담/둘째 아담 등). 핵심 성구 인용.",
  "life_q2": "나눔 질문 2 — '이번 한 주 동안...' 형태 + 구체적 행동 범주를 괄호로 예시(말씀·기도·나눔·내려놓음 등).",
  "s4": [
    "기도제목 1 — 간결한 한 문장, '~하게 하소서/하옵소서' 형태",
    "기도제목 2 — 간결한 한 문장"
  ]
}\
"""

BULLET = ''   # HWP 특수 기호
TAB    = '\t'

# ── 출력 형식 상수 ──
FMT_TXT  = 'txt'
FMT_DOCX = 'docx'
FMT_MD   = 'md'
FMT_HWPX = 'hwpx'
FMT_EXT  = {FMT_TXT: '.txt', FMT_DOCX: '.docx', FMT_MD: '.md', FMT_HWPX: '.hwpx'}


def build_txt(g: dict, doc_title: str = '') -> str:
    lines = []
    if doc_title.strip():
        lines.append(doc_title.strip())
    lines += [
        f"{g['title']}  ({g['scripture']})",
        TAB,
        f"{BULLET} 신앙고백: 사도신경{TAB}{BULLET} 찬송: {g['hymn_open']}장",
        TAB,
        '1. 마을예배 속으로',
    ]
    for i, q in enumerate(g['s1']):
        lines.append(f"{i+1})  {q}")
    lines += [
        TAB,
        '2. 말씀 속으로(본문을 한 번은 함께 읽고, 한 번은 각자 묵상하며 읽어 봅시다.)',
        TAB,
        '3. 삶 속으로',
        TAB,
        g['life_sum1'],
        TAB,
        f"1)  {g['life_q1']}",
        TAB,
        g['life_sum2'],
        TAB,
        f"2)  {g['life_q2']}",
        TAB,
        '4. 합심기도(기도 후, 마을장이 마무리)',
    ]
    for i, p in enumerate(g['s4']):
        lines.append(f"{i+1})  {p}")
    lines += [
        '3)  서로의 기도제목을 나누고 중보하며 기도합시다.',
        '4)  교회를 위해서 기도합시다.',
        TAB,
        f"{BULLET} 찬송: {g['hymn_close']}장{TAB}{BULLET} 주기도문: 다같이",
    ]
    return '\n'.join(lines)


def build_md(g: dict, doc_title: str = '') -> str:
    lines = []
    if doc_title.strip():
        lines += [f"# {doc_title.strip()}", '']
    lines += [
        f"## {g['title']}  ({g['scripture']})",
        '',
        '---',
        '',
        f"> ● 신앙고백: 사도신경　　● 찬송: **{g['hymn_open']}장**",
        '',
        '---',
        '',
        '### 1. 마을예배 속으로',
        '',
    ]
    for q in g['s1']:
        lines.append(f"- {q}")
    lines += [
        '',
        '---',
        '',
        '### 2. 말씀 속으로',
        '',
        '*본문을 한 번은 함께 읽고, 한 번은 각자 묵상하며 읽어 봅시다.*',
        '',
        '---',
        '',
        '### 3. 삶 속으로',
        '',
        g['life_sum1'],
        '',
        f"> **1) {g['life_q1']}**",
        '',
        g['life_sum2'],
        '',
        f"> **2) {g['life_q2']}**",
        '',
        '---',
        '',
        '### 4. 합심기도 *(기도 후, 마을장이 마무리)*',
        '',
    ]
    for p in g['s4']:
        lines.append(f"- {p}")
    lines += [
        '- 서로의 기도제목을 나누고 중보하며 기도합시다.',
        '- 교회를 위해서 기도합시다.',
        '',
        '---',
        '',
        f"> ● 찬송: **{g['hymn_close']}장**　　● 주기도문: 다같이",
    ]
    return '\n'.join(lines)


def build_docx(g: dict, doc_title: str = '') -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx 패키지가 없습니다.\npip install python-docx")

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin  = Cm(3.0)
    sec.right_margin = Cm(3.0)
    sec.top_margin   = Cm(2.5)
    sec.bottom_margin = Cm(2.5)

    def p(text='', size=10.5, bold=False, italic=False,
          center=False, color=None, sb=0, sa=4):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(sb)
        para.paragraph_format.space_after  = Pt(sa)
        para.paragraph_format.left_indent  = Pt(0)
        para.alignment = (WD_ALIGN_PARAGRAPH.CENTER
                          if center else WD_ALIGN_PARAGRAPH.JUSTIFY)
        if text:
            run = para.add_run(text)
            run.font.name   = '맑은 고딕'
            run.font.size   = Pt(size)
            run.font.bold   = bold
            run.font.italic = italic
            if color:
                run.font.color.rgb = RGBColor(*color)
        return para

    def divider():
        para = doc.add_paragraph('─' * 46)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after  = Pt(6)
        run = para.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)

    # 제목
    if doc_title.strip():
        p(doc_title.strip(), size=16, bold=True, center=True,
          color=(26, 26, 46), sb=0, sa=4)
    p(f"{g['title']}  ({g['scripture']})", size=13, bold=True, center=True,
      color=(22, 33, 62), sb=0, sa=6)
    divider()
    p(f"● 신앙고백: 사도신경    ● 찬송: {g['hymn_open']}장",
      size=10.5, center=True, color=(80, 80, 80), sb=0, sa=12)

    # 1. 마을예배 속으로
    p('1. 마을예배 속으로', size=11, bold=True, color=(15, 52, 96), sb=4, sa=5)
    for i, q in enumerate(g['s1']):
        p(f"  {i+1})  {q}", size=10.5, sb=0, sa=3)

    # 2. 말씀 속으로
    p('2. 말씀 속으로', size=11, bold=True, color=(15, 52, 96), sb=10, sa=3)
    p('(본문을 한 번은 함께 읽고, 한 번은 각자 묵상하며 읽어 봅시다.)',
      size=9.5, italic=True, color=(100, 100, 100), sb=0, sa=12)

    # 3. 삶 속으로
    p('3. 삶 속으로', size=11, bold=True, color=(15, 52, 96), sb=4, sa=6)
    p(g['life_sum1'], size=10.5, sb=0, sa=6)
    p(f"1)  {g['life_q1']}", size=10.5, bold=True, sb=2, sa=10)
    p(g['life_sum2'], size=10.5, sb=0, sa=6)
    p(f"2)  {g['life_q2']}", size=10.5, bold=True, sb=2, sa=12)

    # 4. 합심기도
    p('4. 합심기도  (기도 후, 마을장이 마무리)',
      size=11, bold=True, color=(15, 52, 96), sb=4, sa=5)
    for i, pt in enumerate(g['s4']):
        p(f"  {i+1})  {pt}", size=10.5, sb=0, sa=3)
    p('  3)  서로의 기도제목을 나누고 중보하며 기도합시다.', size=10.5, sb=0, sa=3)
    p('  4)  교회를 위해서 기도합시다.', size=10.5, sb=0, sa=10)

    divider()
    p(f"● 찬송: {g['hymn_close']}장    ● 주기도문: 다같이",
      size=10.5, center=True, color=(80, 80, 80), sb=0, sa=4)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_hwpx(g: dict, doc_title: str = '') -> bytes:
    """최소 HWPX (ZIP + HWPML) 생성 — 한컴오피스 2018+ 에서 열림"""

    def esc(s):
        return (str(s).replace('&', '&amp;').replace('<', '&lt;')
                      .replace('>', '&gt;').replace('"', '&quot;'))

    pid = [0]

    def para(text, bold=False):
        pid[0] += 1
        t = esc(text)
        bold_el = '<hsp:bold>true</hsp:bold>' if bold else ''
        return (
            f'<hsp:p hsp:id="{pid[0]}" hsp:paraPrIDRef="0" hsp:styleIDRef="0" '
            f'hsp:pageBreak="false" hsp:columnBreak="false">'
            f'<hsp:run hsp:charPrIDRef="0">'
            + (f'<hsp:charPr hsp:id="0">{bold_el}</hsp:charPr>' if bold else '')
            + f'<hsp:t>{t}</hsp:t>'
            f'</hsp:run></hsp:p>\n'
        )

    paras = []
    if doc_title.strip():
        paras.append(para(doc_title.strip(), bold=True))
    paras.append(para(f"{g['title']}  ({g['scripture']})", bold=True))
    paras.append(para(''))
    paras.append(para(f"● 신앙고백: 사도신경  ● 찬송: {g['hymn_open']}장"))
    paras.append(para(''))
    paras.append(para('1. 마을예배 속으로', bold=True))
    for i, q in enumerate(g['s1']):
        paras.append(para(f"{i+1})  {q}"))
    paras.append(para(''))
    paras.append(para('2. 말씀 속으로(본문을 한 번은 함께 읽고, 한 번은 각자 묵상하며 읽어 봅시다.)', bold=True))
    paras.append(para(''))
    paras.append(para('3. 삶 속으로', bold=True))
    paras.append(para(''))
    paras.append(para(g['life_sum1']))
    paras.append(para(''))
    paras.append(para(f"1)  {g['life_q1']}"))
    paras.append(para(''))
    paras.append(para(g['life_sum2']))
    paras.append(para(''))
    paras.append(para(f"2)  {g['life_q2']}"))
    paras.append(para(''))
    paras.append(para('4. 합심기도(기도 후, 마을장이 마무리)', bold=True))
    for i, p in enumerate(g['s4']):
        paras.append(para(f"{i+1})  {p}"))
    paras.append(para('3)  서로의 기도제목을 나누고 중보하며 기도합시다.'))
    paras.append(para('4)  교회를 위해서 기도합시다.'))
    paras.append(para(''))
    paras.append(para(f"● 찬송: {g['hymn_close']}장  ● 주기도문: 다같이"))

    ns = 'http://www.hancom.co.kr/hwpml/2012/section'
    section_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        f'<hsp:sec xmlns:hsp="{ns}" hsp:id="0">\n'
        + ''.join(paras)
        + '</hsp:sec>\n'
    ).encode('utf-8')

    container_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="no"?>\n'
        '<container>\n'
        '  <rootfiles>\n'
        '    <rootfile full-path="contents.hpf" media-type="application/x-hwpml+xml"/>\n'
        '  </rootfiles>\n'
        '</container>\n'
    ).encode('utf-8')

    hpf_xml = (
        '<?xml version="1.0" encoding="UTF-8"?>\n'
        '<opf:package xmlns:opf="http://www.idpf.org/2007/opf" version="2.0">\n'
        '  <opf:metadata>\n'
        '    <dc:title xmlns:dc="http://purl.org/dc/elements/1.1/">소그룹교안</dc:title>\n'
        '  </opf:metadata>\n'
        '  <opf:manifest>\n'
        '    <opf:item id="sec0" href="Contents/section0.xml"'
        ' media-type="application/x-hwpml+xml"/>\n'
        '  </opf:manifest>\n'
        '  <opf:spine>\n'
        '    <opf:itemref idref="sec0"/>\n'
        '  </opf:spine>\n'
        '</opf:package>\n'
    ).encode('utf-8')

    buf = BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        mi = zipfile.ZipInfo('mimetype')
        mi.compress_type = zipfile.ZIP_STORED
        zf.writestr(mi, 'application/x-hwp+zip')
        zf.writestr('META-INF/container.xml', container_xml)
        zf.writestr('contents.hpf', hpf_xml)
        zf.writestr('Contents/section0.xml', section_xml)
    return buf.getvalue()


def _parse_json(raw: str) -> dict:
    if '```' in raw:
        raw = re.sub(r'```(?:json)?', '', raw).strip()
    m = re.search(r'\{.*\}', raw, re.DOTALL)
    if m:
        raw = m.group(0)
    return json.loads(raw)


# ──────────────────────────── 백엔드 ─────────────────────────────────────────

BACKEND_CLAUDE    = 'claude_cli'
BACKEND_CODEX     = 'codex_cli'
BACKEND_ANTHROPIC = 'api_anthropic'
BACKEND_GEMINI    = 'api_gemini'
BACKEND_CHATGPT   = 'api_chatgpt'

BACKEND_LABELS = {
    BACKEND_CLAUDE:    'Claude CLI',
    BACKEND_CODEX:     'Codex CLI (OpenAI)',
    BACKEND_ANTHROPIC: 'Anthropic API',
    BACKEND_GEMINI:    'Gemini API (Google)',
    BACKEND_CHATGPT:   'ChatGPT API (OpenAI)',
}

BACKEND_MODELS = {
    BACKEND_CLAUDE: [
        ('sonnet-4-6 (기본)', 'claude-sonnet-4-6'),
        ('opus-4-8 (고품질)', 'claude-opus-4-8'),
        ('haiku-4-5 (빠름)',  'claude-haiku-4-5-20251001'),
    ],
    BACKEND_CODEX: [
        ('gpt-5.5 (ChatGPT 기본)', 'gpt-5.5'),
        ('o4-mini (API 계정)',      'o4-mini'),
        ('gpt-4o (API 계정)',       'gpt-4o'),
    ],
    BACKEND_ANTHROPIC: [
        ('sonnet-4-6 (기본)', 'claude-sonnet-4-6'),
        ('opus-4-8 (고품질)', 'claude-opus-4-8'),
        ('haiku-4-5 (빠름)',  'claude-haiku-4-5-20251001'),
    ],
    BACKEND_GEMINI: [
        ('gemini-2.0-flash (기본)', 'gemini-2.0-flash'),
        ('gemini-2.5-pro (고품질)', 'gemini-2.5-pro'),
        ('gemini-1.5-pro',          'gemini-1.5-pro'),
    ],
    BACKEND_CHATGPT: [
        ('gpt-4o (기본)',      'gpt-4o'),
        ('gpt-4o-mini (빠름)', 'gpt-4o-mini'),
        ('o3',                 'o3'),
    ],
}


def find_claude_exe() -> str:
    candidates = [
        os.environ.get('CLAUDE_CODE_EXECPATH', ''),
        str(Path.home() / '.local' / 'bin' / 'claude.exe'),
        str(Path.home() / '.local' / 'bin' / 'claude'),
        r'C:\Users\Owner\.local\bin\claude.exe',
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    raise FileNotFoundError(
        "claude.exe를 찾을 수 없습니다.\nClaude Code CLI를 설치해주세요."
    )


def find_codex_exe() -> str:
    npm_dir = str(Path.home() / 'AppData' / 'Roaming' / 'npm')
    candidates = [
        os.environ.get('CODEX_PATH', ''),
        os.path.join(npm_dir, 'codex.cmd'),
        os.path.join(npm_dir, 'codex.ps1'),
        os.path.join(npm_dir, 'codex'),
        '/usr/local/bin/codex',
        str(Path.home() / '.local' / 'bin' / 'codex'),
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    try:
        r = subprocess.run(['where.exe', 'codex.cmd'], capture_output=True, text=True,
                           encoding='utf-8', errors='replace', timeout=5)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout.strip().splitlines()[0].strip()
    except Exception:
        pass
    raise FileNotFoundError(
        "codex를 찾을 수 없습니다.\nnpm install -g @openai/codex 로 설치해주세요."
    )


def _codex_base_cmd(exe: str) -> list:
    ext = Path(exe).suffix.lower()
    if ext == '.cmd':
        return ['cmd', '/c', exe]
    if ext == '.ps1':
        return ['powershell', '-ExecutionPolicy', 'Bypass', '-File', exe]
    return [exe]


def detect_available_backends() -> list:
    available = []
    try:
        find_claude_exe()
        available.append(BACKEND_CLAUDE)
    except FileNotFoundError:
        pass
    try:
        find_codex_exe()
        available.append(BACKEND_CODEX)
    except FileNotFoundError:
        pass
    available += [BACKEND_ANTHROPIC, BACKEND_GEMINI, BACKEND_CHATGPT]
    return available


def generate_via_claude_cli(sermon: str, model: str) -> dict:
    exe    = find_claude_exe()
    prompt = f"{SYSTEM_PROMPT}\n\n다음 설교문으로 교안을 작성해주세요:\n\n{sermon}"
    r = subprocess.run(
        [exe, '-p', prompt, '--output-format', 'text', '--model', model],
        capture_output=True, text=True, encoding='utf-8', errors='replace',
        timeout=180, creationflags=_NO_WINDOW,
    )
    if r.returncode != 0:
        raise RuntimeError(f"Claude CLI 오류:\n{r.stderr[-600:]}")
    return _parse_json(r.stdout.strip())


def generate_via_codex_cli(sermon: str, model: str) -> dict:
    exe    = find_codex_exe()
    prompt = f"{SYSTEM_PROMPT}\n\n다음 설교문으로 교안을 작성해주세요:\n\n{sermon}"
    cmd = _codex_base_cmd(exe) + ['exec', '--skip-git-repo-check', '-m', model, '-']
    r = subprocess.run(
        cmd, input=prompt,
        capture_output=True, text=True, encoding='utf-8', errors='replace',
        timeout=180, creationflags=_NO_WINDOW,
    )
    if r.returncode != 0:
        err_msg = r.stderr
        m = re.search(r'ERROR:\s*(\{.*?\})', err_msg)
        if m:
            try:
                d = json.loads(m.group(1))
                err_msg = d.get('error', {}).get('message', err_msg[-600:])
            except Exception:
                err_msg = err_msg[-600:]
        else:
            err_msg = err_msg[-600:]
        raise RuntimeError(f"Codex CLI 오류 (code {r.returncode}):\n{err_msg}")
    return _parse_json(r.stdout.strip())


def generate_via_anthropic(api_key: str, sermon: str, model: str) -> dict:
    try:
        import anthropic
    except ImportError:
        raise ImportError("anthropic 패키지가 없습니다.\npip install anthropic")
    client = anthropic.Anthropic(api_key=api_key)
    msg = client.messages.create(
        model=model, max_tokens=4096, system=SYSTEM_PROMPT,
        messages=[{'role': 'user',
                   'content': f'다음 설교문으로 교안을 작성해주세요:\n\n{sermon}'}],
    )
    return _parse_json(msg.content[0].text.strip())


def generate_via_gemini(api_key: str, sermon: str, model: str) -> dict:
    try:
        import google.generativeai as genai
    except ImportError:
        raise ImportError(
            "google-generativeai 패키지가 없습니다.\npip install google-generativeai"
        )
    genai.configure(api_key=api_key)
    m = genai.GenerativeModel(model, system_instruction=SYSTEM_PROMPT)
    resp = m.generate_content(f'다음 설교문으로 교안을 작성해주세요:\n\n{sermon}')
    return _parse_json(resp.text.strip())


def generate_via_chatgpt(api_key: str, sermon: str, model: str) -> dict:
    try:
        from openai import OpenAI
    except ImportError:
        raise ImportError("openai 패키지가 없습니다.\npip install openai")
    client = OpenAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model, max_tokens=4096,
        messages=[
            {'role': 'system', 'content': SYSTEM_PROMPT},
            {'role': 'user',   'content': f'다음 설교문으로 교안을 작성해주세요:\n\n{sermon}'},
        ],
    )
    return _parse_json(resp.choices[0].message.content.strip())


# ──────────────────────────── CLI 설치 도우미 ─────────────────────────────────

def _check_npm() -> tuple:
    try:
        r = subprocess.run(['npm', '--version'], capture_output=True, text=True,
                           encoding='utf-8', errors='replace', timeout=5)
        if r.returncode == 0:
            return True, r.stdout.strip()
    except Exception:
        pass
    return False, ''


class InstallDialog(tk.Toplevel):
    """Claude CLI / Codex CLI 설치 도우미 다이얼로그"""

    def __init__(self, parent):
        super().__init__(parent)
        self.title('AI 엔진 설치 도우미')
        self.geometry('620x460')
        self.resizable(False, False)
        self.grab_set()
        self._build()
        self._refresh()

    def _build(self):
        P = dict(padx=12, pady=6)

        ttk.Label(self, text='감지된 AI 엔진 및 설치 상태',
                  font=('맑은 고딕', 11, 'bold')).pack(**P, anchor='w')

        # npm 상태
        frm_npm = ttk.LabelFrame(self, text=' Node.js / npm ')
        frm_npm.pack(fill='x', **P)
        self.lbl_npm = ttk.Label(frm_npm, text='확인 중...')
        self.lbl_npm.pack(side='left', padx=8, pady=4)
        ttk.Button(frm_npm, text='Node.js 다운로드',
                   command=lambda: os.startfile('https://nodejs.org')).pack(
                       side='right', padx=8)

        # Claude CLI
        frm_cl = ttk.LabelFrame(self, text=' Claude CLI (Anthropic) ')
        frm_cl.pack(fill='x', **P)
        self.lbl_cl = ttk.Label(frm_cl, text='확인 중...')
        self.lbl_cl.pack(side='left', padx=8, pady=4, fill='x', expand=True)
        self.btn_cl = ttk.Button(frm_cl, text='npm 설치',
                                 command=self._install_claude)
        self.btn_cl.pack(side='right', padx=8)

        # Codex CLI
        frm_cx = ttk.LabelFrame(self, text=' Codex CLI (OpenAI) ')
        frm_cx.pack(fill='x', **P)
        self.lbl_cx = ttk.Label(frm_cx, text='확인 중...')
        self.lbl_cx.pack(side='left', padx=8, pady=4, fill='x', expand=True)
        self.btn_cx = ttk.Button(frm_cx, text='npm 설치',
                                 command=self._install_codex)
        self.btn_cx.pack(side='right', padx=8)

        # 로그
        frm_log = ttk.LabelFrame(self, text=' 설치 로그 ')
        frm_log.pack(fill='both', expand=True, **P)
        self.log = scrolledtext.ScrolledText(frm_log, height=8, wrap='word',
                                             font=('Consolas', 9))
        self.log.pack(fill='both', expand=True, padx=4, pady=4)

        ttk.Button(self, text='새로고침', command=self._refresh).pack(side='left', padx=12, pady=8)
        ttk.Button(self, text='닫기', command=self.destroy).pack(side='right', padx=12, pady=8)

    def _log(self, msg: str):
        self.log.insert('end', msg + '\n')
        self.log.see('end')

    def _refresh(self):
        npm_ok, npm_ver = _check_npm()
        if npm_ok:
            self.lbl_npm.config(text=f'설치됨: npm v{npm_ver}', foreground='green')
        else:
            self.lbl_npm.config(text='npm을 찾을 수 없습니다. Node.js를 먼저 설치하세요.',
                                foreground='red')

        try:
            exe = find_claude_exe()
            self.lbl_cl.config(text=f'연결됨: {exe}', foreground='green')
            self.btn_cl.config(state='disabled')
        except FileNotFoundError:
            self.lbl_cl.config(text='설치되지 않음', foreground='red')
            self.btn_cl.config(state='normal' if npm_ok else 'disabled')

        try:
            exe = find_codex_exe()
            self.lbl_cx.config(text=f'연결됨: {exe}', foreground='green')
            self.btn_cx.config(state='disabled')
        except FileNotFoundError:
            self.lbl_cx.config(text='설치되지 않음', foreground='red')
            self.btn_cx.config(state='normal' if npm_ok else 'disabled')

    def _run_npm_install(self, pkg: str, btn: ttk.Button):
        btn.config(state='disabled')
        self._log(f'> npm install -g {pkg}')

        def worker():
            r = subprocess.run(
                ['npm', 'install', '-g', pkg],
                capture_output=True, text=True, encoding='utf-8', errors='replace',
                timeout=120, creationflags=_NO_WINDOW,
            )
            out = (r.stdout + r.stderr).strip()
            self.after(0, lambda: self._log(out or '(출력 없음)'))
            if r.returncode == 0:
                self.after(0, lambda: self._log('✓ 설치 완료'))
            else:
                self.after(0, lambda: self._log(f'✗ 오류 (code {r.returncode})'))
            self.after(0, self._refresh)

        threading.Thread(target=worker, daemon=True).start()

    def _install_claude(self):
        self._run_npm_install('@anthropic-ai/claude-code', self.btn_cl)

    def _install_codex(self):
        self._run_npm_install('@openai/codex', self.btn_cx)


# ──────────────────────────── 설정 ───────────────────────────────────────────

CFG_PATH = Path.home() / '.village_worship_cfg.json'


def _default_output_dir() -> str:
    docs = Path.home() / 'Documents'
    docs.mkdir(exist_ok=True)
    return str(docs)


def load_cfg() -> dict:
    if CFG_PATH.exists():
        try:
            data = json.loads(CFG_PATH.read_text(encoding='utf-8'))
            # v1.2 → v1.3+ 마이그레이션
            if 'api_key' in data and 'api_anthropic' not in data:
                data['api_anthropic'] = data.pop('api_key', '')
            if data.get('backend') == 'api_key':
                data['backend'] = BACKEND_ANTHROPIC
            data.setdefault('output_dir', _default_output_dir())
            data.setdefault('doc_title', '금주의 마을예배')
            data.setdefault('output_fmt', FMT_TXT)
            return data
        except Exception:
            pass
    return {
        'backend':       BACKEND_CLAUDE,
        'model':         'claude-sonnet-4-6',
        'output_dir':    _default_output_dir(),
        'doc_title':     '금주의 마을예배',
        'output_fmt':    FMT_TXT,
        'api_anthropic': '',
        'api_gemini':    '',
        'api_chatgpt':   '',
    }


def save_cfg(cfg: dict):
    CFG_PATH.write_text(json.dumps(cfg, ensure_ascii=False, indent=2), encoding='utf-8')


# ──────────────────────────── GUI ─────────────────────────────────────────────

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('소그룹교안 생성기 v1.4')
        self.geometry('980x820')
        self.minsize(760, 620)
        self.resizable(True, True)
        self.cfg           = load_cfg()
        self._data         = None   # 마지막 생성 데이터 (포맷 변환용)
        self._result_txt   = ''
        self._progress_id  = None
        self.v_key_anthropic = tk.StringVar(value=self.cfg.get('api_anthropic', ''))
        self.v_key_gemini    = tk.StringVar(value=self.cfg.get('api_gemini', ''))
        self.v_key_chatgpt   = tk.StringVar(value=self.cfg.get('api_chatgpt', ''))
        self._available = detect_available_backends()
        if self.cfg.get('backend') not in self._available:
            self.cfg['backend'] = self._available[0]
        self._build_ui()
        self._setup_dnd()
        # 저장 폴더 미설정 시에만 안내 (기본값 Documents가 이미 있으므로 거의 안 뜸)
        if not self.cfg.get('output_dir') or not Path(self.cfg['output_dir']).exists():
            self.cfg['output_dir'] = _default_output_dir()
            self.after(100, lambda: self.v_outdir.set(self.cfg['output_dir']))

    # ── UI 구성 ──────────────────────────────────────────────────────────────

    def _build_ui(self):
        P = dict(padx=10, pady=4)

        # ── AI 엔진 선택 ──
        frm_engine = ttk.LabelFrame(self, text=' AI 엔진 선택 ')
        frm_engine.pack(fill='x', **P)
        top = ttk.Frame(frm_engine)
        top.pack(fill='x', padx=6, pady=(6, 2))
        self.v_backend = tk.StringVar(value=self.cfg.get('backend', BACKEND_CLAUDE))
        for b in self._available:
            ttk.Radiobutton(top, text=BACKEND_LABELS[b], variable=self.v_backend,
                            value=b, command=self._on_backend_change).pack(side='left', padx=8)
        ttk.Button(top, text='설치 도우미', width=10,
                   command=lambda: InstallDialog(self)).pack(side='right', padx=6)
        self.frm_sub = ttk.Frame(frm_engine)
        self.frm_sub.pack(fill='x', padx=6, pady=(2, 6))
        self._rebuild_sub()

        # ── 모델 선택 ──
        self.frm_model = ttk.LabelFrame(self, text=' 모델 ')
        self.frm_model.pack(fill='x', **P)
        self.v_model = tk.StringVar(value=self.cfg.get('model', 'claude-sonnet-4-6'))
        self._rebuild_models()

        # ── 교안 상단 제목 + 출력 형식 ──
        frm_opt = ttk.Frame(self)
        frm_opt.pack(fill='x', **P)

        frm_title = ttk.LabelFrame(frm_opt, text=' 교안 상단 제목  (빈칸 = 제목줄 생략) ')
        frm_title.pack(side='left', fill='x', expand=True, padx=(0, 6))
        self.v_doc_title = tk.StringVar(value=self.cfg.get('doc_title', '금주의 마을예배'))
        ttk.Entry(frm_title, textvariable=self.v_doc_title, width=28).pack(
            side='left', padx=6, pady=5)
        ttk.Label(frm_title, text='예) 금주의 마을예배 / 목장 / 셀 / 구역',
                  foreground='gray').pack(side='left', padx=4)

        frm_fmt = ttk.LabelFrame(frm_opt, text=' 출력 형식 ')
        frm_fmt.pack(side='right')
        self.v_fmt = tk.StringVar(value=self.cfg.get('output_fmt', FMT_TXT))
        for label, val in [('TXT', FMT_TXT), ('DOCX', FMT_DOCX), ('MD', FMT_MD)]:
            ttk.Radiobutton(frm_fmt, text=label, variable=self.v_fmt,
                            value=val, command=self._on_fmt_change).pack(
                                side='left', padx=8, pady=5)

        # ── 설교 파일 ──
        frm = ttk.LabelFrame(self, text=' 설교 파일  (PDF · DOCX · HWPX · TXT) ')
        frm.pack(fill='x', **P)
        self.v_in = tk.StringVar()
        self._ent_in = ttk.Entry(frm, textvariable=self.v_in, width=72)
        self._ent_in.pack(side='left', padx=6, pady=5, fill='x', expand=True)
        ttk.Button(frm, text='찾아보기', command=self._browse_in).pack(side='left', padx=4)

        # ── 저장 폴더 ──
        frm = ttk.LabelFrame(self, text=' 저장 폴더 ')
        frm.pack(fill='x', **P)
        self.v_outdir = tk.StringVar(value=self.cfg.get('output_dir', ''))
        ttk.Entry(frm, textvariable=self.v_outdir, width=72).pack(
            side='left', padx=6, pady=5, fill='x', expand=True)
        ttk.Button(frm, text='찾아보기', command=self._browse_outdir).pack(side='left', padx=4)

        # ── 날짜 + 파일명 ──
        frm = ttk.Frame(self)
        frm.pack(fill='x', padx=10, pady=3)
        ttk.Label(frm, text='예배 날짜:').pack(side='left')
        self.v_date = tk.StringVar(value=datetime.now().strftime('%Y%m%d'))
        ttk.Entry(frm, textvariable=self.v_date, width=10).pack(side='left', padx=4)
        ttk.Label(frm, text='  파일명:').pack(side='left', padx=(8, 0))
        self.v_outname = tk.StringVar(value='YYYYMMDD_소그룹교안_제목.txt')
        ttk.Entry(frm, textvariable=self.v_outname, width=48).pack(
            side='left', padx=4, fill='x', expand=True)

        # ── 버튼 + 상태 ──
        frm = ttk.Frame(self)
        frm.pack(fill='x', padx=10, pady=8)
        self.btn_gen = ttk.Button(frm, text='교안 생성하기', width=16, command=self._run)
        self.btn_gen.pack(side='left', ipadx=6, ipady=5)
        self.btn_save = ttk.Button(frm, text='파일 저장', width=12,
                                   command=self._save, state='disabled')
        self.btn_save.pack(side='left', padx=8, ipadx=6, ipady=5)
        self.btn_copy = ttk.Button(frm, text='클립보드 복사', width=14,
                                   command=self._copy, state='disabled')
        self.btn_copy.pack(side='left', ipadx=6, ipady=5)
        self.lbl_st = ttk.Label(frm, text='설교 파일을 선택하고 "교안 생성하기"를 클릭하세요.',
                                foreground='gray')
        self.lbl_st.pack(side='left', padx=10)

        # ── 미리보기 ──
        frm = ttk.LabelFrame(self, text=' 미리보기  (TXT 모드에서 직접 수정 가능) ')
        frm.pack(fill='both', expand=True, padx=10, pady=4)
        self.txt = scrolledtext.ScrolledText(
            frm, wrap='word', font=('맑은 고딕', 10), undo=True)
        self.txt.pack(fill='both', expand=True, padx=4, pady=4)
        self._setup_preview_tags()

    # ── 동적 패널 ────────────────────────────────────────────────────────────

    def _rebuild_sub(self):
        for w in self.frm_sub.winfo_children():
            w.destroy()
        backend = self.v_backend.get()
        if backend in (BACKEND_CLAUDE, BACKEND_CODEX):
            try:
                exe = find_claude_exe() if backend == BACKEND_CLAUDE else find_codex_exe()
                text, color = f'연결됨: {exe}', 'green'
            except FileNotFoundError as e:
                text, color = str(e).split('\n')[0], 'red'
            ttk.Label(self.frm_sub, text=text, foreground=color).pack(side='left', pady=2)
        elif backend in (BACKEND_ANTHROPIC, BACKEND_GEMINI, BACKEND_CHATGPT):
            var_map = {
                BACKEND_ANTHROPIC: (self.v_key_anthropic, 'Anthropic'),
                BACKEND_GEMINI:    (self.v_key_gemini,    'Gemini'),
                BACKEND_CHATGPT:   (self.v_key_chatgpt,   'ChatGPT'),
            }
            var, name = var_map[backend]
            ttk.Label(self.frm_sub, text=f'{name} API Key:').pack(side='left')
            self._ent_key = ttk.Entry(self.frm_sub, textvariable=var, show='*', width=52)
            self._ent_key.pack(side='left', padx=4, fill='x', expand=True)
            ttk.Button(self.frm_sub, text='저장', width=6,
                       command=lambda b=backend: self._save_key(b)).pack(side='left', padx=2)
            ttk.Button(self.frm_sub, text='표시/숨김',
                       command=self._toggle_key).pack(side='left', padx=4)

    def _rebuild_models(self, force_reset: bool = False):
        for w in self.frm_model.winfo_children():
            w.destroy()
        backend = self.v_backend.get()
        models  = BACKEND_MODELS.get(backend, BACKEND_MODELS[BACKEND_CLAUDE])
        valid   = [v for _, v in models]
        if force_reset or self.v_model.get() not in valid:
            self.v_model.set(valid[0])
        for label, val in models:
            ttk.Radiobutton(self.frm_model, text=label,
                            variable=self.v_model, value=val).pack(
                                side='left', padx=12, pady=5)

    def _on_backend_change(self):
        self.cfg['backend'] = self.v_backend.get()
        save_cfg(self.cfg)
        self._rebuild_sub()
        self._rebuild_models(force_reset=True)  # 백엔드 바뀌면 모델 항상 초기화

    def _setup_preview_tags(self):
        self.txt.tag_config('h1', font=('맑은 고딕', 15, 'bold'),
                            justify='center', spacing1=8, spacing3=8,
                            foreground='#1a1a2e')
        self.txt.tag_config('h2', font=('맑은 고딕', 12, 'bold'),
                            justify='center', spacing1=4, spacing3=8,
                            foreground='#16213e')
        self.txt.tag_config('section', font=('맑은 고딕', 11, 'bold'),
                            spacing1=12, spacing3=5, foreground='#0f3460')
        self.txt.tag_config('body', font=('맑은 고딕', 10),
                            lmargin1=8, lmargin2=8, spacing1=2, spacing3=2)
        self.txt.tag_config('sub', font=('맑은 고딕', 10),
                            lmargin1=28, lmargin2=28, spacing1=2, spacing3=2)
        self.txt.tag_config('qbold', font=('맑은 고딕', 10, 'bold'),
                            lmargin1=16, lmargin2=16, spacing1=5, spacing3=5)
        self.txt.tag_config('meta', font=('맑은 고딕', 10),
                            foreground='#444444', justify='center', spacing3=6)
        self.txt.tag_config('note', font=('맑은 고딕', 9, 'italic'),
                            foreground='#777777', lmargin1=20, lmargin2=20,
                            spacing1=2, spacing3=8)
        self.txt.tag_config('hr', font=('맑은 고딕', 8),
                            foreground='#bbbbbb', justify='center',
                            spacing1=4, spacing3=6)

    def _render_preview(self):
        if self._data is None:
            return
        fmt       = self.v_fmt.get()
        doc_title = self.v_doc_title.get()
        data      = self._data

        self.txt.delete('1.0', 'end')

        if fmt == FMT_TXT:
            self.txt.insert('1.0', build_txt(data, doc_title))
            return

        # DOCX / MD: 서식 태그 미리보기
        def ins(text, tag=None):
            pos = self.txt.index('end')
            self.txt.insert('end', text)
            if tag:
                self.txt.tag_add(tag, pos, 'end')

        HR = '─' * 50 + '\n'

        if doc_title.strip():
            ins(doc_title.strip() + '\n', 'h1')
        ins(f"{data['title']}  ({data['scripture']})\n", 'h2')
        ins(HR, 'hr')
        ins(f"● 신앙고백: 사도신경    ● 찬송: {data['hymn_open']}장\n", 'meta')
        ins(HR, 'hr')
        ins('\n')
        ins('1. 마을예배 속으로\n', 'section')
        for i, q in enumerate(data['s1']):
            ins(f"  {i+1})  {q}\n", 'sub')
        ins('\n')
        ins(HR, 'hr')
        ins('\n')
        ins('2. 말씀 속으로\n', 'section')
        ins('  본문을 한 번은 함께 읽고, 한 번은 각자 묵상하며 읽어 봅시다.\n', 'note')
        ins('\n')
        ins(HR, 'hr')
        ins('\n')
        ins('3. 삶 속으로\n', 'section')
        ins('\n')
        ins(data['life_sum1'] + '\n', 'body')
        ins('\n')
        ins(f"  1)  {data['life_q1']}\n", 'qbold')
        ins('\n')
        ins(data['life_sum2'] + '\n', 'body')
        ins('\n')
        ins(f"  2)  {data['life_q2']}\n", 'qbold')
        ins('\n')
        ins(HR, 'hr')
        ins('\n')
        ins('4. 합심기도  (기도 후, 마을장이 마무리)\n', 'section')
        for i, pr in enumerate(data['s4']):
            ins(f"  {i+1})  {pr}\n", 'sub')
        ins('  3)  서로의 기도제목을 나누고 중보하며 기도합시다.\n', 'sub')
        ins('  4)  교회를 위해서 기도합시다.\n', 'sub')
        ins('\n')
        ins(HR, 'hr')
        ins(f"● 찬송: {data['hymn_close']}장    ● 주기도문: 다같이\n", 'meta')

    def _on_fmt_change(self):
        fmt = self.v_fmt.get()
        self.cfg['output_fmt'] = fmt
        save_cfg(self.cfg)
        cur = self.v_outname.get()
        for ext in FMT_EXT.values():
            if cur.endswith(ext):
                self.v_outname.set(cur[:-len(ext)] + FMT_EXT[fmt])
                break
        self._render_preview()

    # ── 최초 실행 ─────────────────────────────────────────────────────────────

    def _prompt_first_run(self):
        messagebox.showinfo(
            '저장 폴더 선택',
            '교안이 저장될 폴더를 선택해주세요.\n(나중에 "저장 폴더" 항목에서 변경할 수 있습니다.)',
        )
        d = filedialog.askdirectory(title='교안 저장 폴더 선택')
        if d:
            self.v_outdir.set(d)
            self.cfg['output_dir'] = d
            save_cfg(self.cfg)

    # ── DnD ──────────────────────────────────────────────────────────────────

    def _setup_dnd(self):
        try:
            from tkinterdnd2 import DND_FILES
            self._ent_in.drop_target_register(DND_FILES)
            self._ent_in.dnd_bind('<<Drop>>', self._on_drop)
        except Exception:
            pass

    def _on_drop(self, event):
        self.v_in.set(event.data.strip('{}'))
        self._suggest_filename(self.v_in.get())

    # ── 이벤트 ───────────────────────────────────────────────────────────────

    def _toggle_key(self):
        e = getattr(self, '_ent_key', None)
        if e:
            e.config(show='' if e.cget('show') else '*')

    def _save_key(self, backend: str):
        var_cfg = {
            BACKEND_ANTHROPIC: (self.v_key_anthropic, 'api_anthropic'),
            BACKEND_GEMINI:    (self.v_key_gemini,    'api_gemini'),
            BACKEND_CHATGPT:   (self.v_key_chatgpt,   'api_chatgpt'),
        }
        var, cfg_key = var_cfg[backend]
        key = var.get().strip()
        if not key:
            messagebox.showwarning('입력 필요', 'API Key를 입력하세요.')
            return
        self.cfg[cfg_key] = key
        save_cfg(self.cfg)
        messagebox.showinfo('저장', 'API Key가 저장되었습니다.')

    def _browse_in(self):
        p = filedialog.askopenfilename(
            title='설교 파일 선택',
            filetypes=[('지원 파일', '*.pdf *.docx *.hwpx *.txt'),
                       ('PDF', '*.pdf'), ('Word', '*.docx'),
                       ('한컴 HWPX', '*.hwpx'), ('텍스트', '*.txt'),
                       ('모든 파일', '*.*')],
        )
        if p:
            self.v_in.set(p)
            self._suggest_filename(p)

    def _suggest_filename(self, path: str):
        stem   = re.sub(r'[\s\+\-\(\)\[\]]+', '', Path(path).stem)
        tlabel = self.v_doc_title.get().strip().replace(' ', '') or '소그룹교안'
        ext    = FMT_EXT.get(self.v_fmt.get(), '.txt')
        self.v_outname.set(f"{self.v_date.get()}_{tlabel}_{stem}{ext}")

    def _browse_outdir(self):
        d = filedialog.askdirectory(title='저장 폴더 선택')
        if d:
            self.v_outdir.set(d)
            self.cfg['output_dir'] = d
            save_cfg(self.cfg)

    # ── 진행 타이머 ───────────────────────────────────────────────────────────

    def _start_progress(self, backend_name: str, sermon_len: int):
        self._gen_start  = datetime.now()
        self._gen_label  = backend_name
        self._gen_slen   = sermon_len
        self._progress_id = None
        self._tick_progress()

    def _tick_progress(self):
        elapsed = int((datetime.now() - self._gen_start).total_seconds())
        self.lbl_st.config(
            text=(f'설교문 {self._gen_slen:,}자 추출 완료.'
                  f' {self._gen_label}가 교안 작성 중... ({elapsed}초 경과)')
        )
        self._progress_id = self.after(1000, self._tick_progress)

    def _stop_progress(self):
        if self._progress_id:
            self.after_cancel(self._progress_id)
            self._progress_id = None

    # ── 생성 ─────────────────────────────────────────────────────────────────

    def _run(self):
        in_path = self.v_in.get().strip()
        if not in_path:
            messagebox.showwarning('입력 필요', '설교 파일을 선택하세요.')
            return
        if not Path(in_path).exists():
            messagebox.showerror('파일 없음', f'파일을 찾을 수 없습니다:\n{in_path}')
            return
        if not self.v_outdir.get().strip():
            messagebox.showwarning('저장 폴더 없음', '저장 폴더를 먼저 선택하세요.')
            self._prompt_first_run()
            return

        backend = self.v_backend.get()
        model   = self.v_model.get()
        api_key = ''
        key_map = {
            BACKEND_ANTHROPIC: (self.v_key_anthropic, 'Anthropic'),
            BACKEND_GEMINI:    (self.v_key_gemini,    'Gemini'),
            BACKEND_CHATGPT:   (self.v_key_chatgpt,   'ChatGPT'),
        }
        if backend in key_map:
            var, name = key_map[backend]
            api_key = var.get().strip()
            if not api_key:
                messagebox.showwarning('API Key 없음', f'{name} API Key를 입력하고 저장하세요.')
                return

        doc_title = self.v_doc_title.get()
        self.cfg.update({'backend': backend, 'model': model,
                         'doc_title': doc_title, 'output_fmt': self.v_fmt.get()})
        save_cfg(self.cfg)

        self.btn_gen.config(state='disabled')
        self.btn_save.config(state='disabled')
        self.btn_copy.config(state='disabled')
        self.txt.delete('1.0', 'end')
        self._set_status('설교 파일 읽는 중...')
        threading.Thread(
            target=self._worker,
            args=(in_path, backend, model, api_key, doc_title),
            daemon=True,
        ).start()

    def _worker(self, in_path, backend, model, api_key, doc_title):
        try:
            sermon = extract_text(in_path)
            if len(sermon.strip()) < 200:
                raise ValueError(
                    f'추출된 텍스트가 너무 짧습니다 ({len(sermon)}자).\n파일을 확인해주세요.')

            bname = BACKEND_LABELS.get(backend, backend)
            self.after(0, lambda: self._start_progress(bname, len(sermon)))

            if backend == BACKEND_CLAUDE:
                data = generate_via_claude_cli(sermon, model)
            elif backend == BACKEND_CODEX:
                data = generate_via_codex_cli(sermon, model)
            elif backend == BACKEND_ANTHROPIC:
                data = generate_via_anthropic(api_key, sermon, model)
            elif backend == BACKEND_GEMINI:
                data = generate_via_gemini(api_key, sermon, model)
            elif backend == BACKEND_CHATGPT:
                data = generate_via_chatgpt(api_key, sermon, model)
            else:
                raise ValueError(f'알 수 없는 백엔드: {backend}')

            self._data      = data
            preview_txt     = build_txt(data, doc_title)
            self._result_txt = preview_txt

            title_clean = re.sub(r'[^\w가-힣]', '', data.get('title', '교안'))[:20]
            tlabel      = doc_title.strip().replace(' ', '') or '소그룹교안'
            ext         = FMT_EXT.get(self.v_fmt.get(), '.txt')
            suggested   = f"{self.v_date.get()}_{tlabel}_{title_clean}{ext}"

            self.after(0, self._stop_progress)
            self.after(0, lambda: self.v_outname.set(suggested))
            self.after(0, self._on_done)

        except Exception as e:
            err = str(e)
            self.after(0, self._stop_progress)
            self.after(0, lambda: messagebox.showerror('오류', err))
            self.after(0, lambda: self._set_status('오류가 발생했습니다.'))
            self.after(0, lambda: self.btn_gen.config(state='normal'))

    def _on_done(self):
        self._render_preview()
        self.btn_gen.config(state='normal')
        self.btn_save.config(state='normal')
        self.btn_copy.config(state='normal')
        self._set_status('교안 생성 완료! 내용 확인 후 "파일 저장"을 클릭하세요.')

    # ── 저장 ─────────────────────────────────────────────────────────────────

    def _save(self):
        if self._data is None:
            return
        fmt       = self.v_fmt.get()
        doc_title = self.v_doc_title.get()
        outname   = self.v_outname.get().strip()
        ext       = FMT_EXT.get(fmt, '.txt')
        if not any(outname.endswith(e) for e in FMT_EXT.values()):
            outname += ext
        out_path = Path(self.v_outdir.get().strip()) / outname
        try:
            out_path.parent.mkdir(parents=True, exist_ok=True)
            if fmt == FMT_TXT:
                # TXT 모드는 미리보기 텍스트 그대로 저장 (직접 수정 반영)
                out_path.write_text(self.txt.get('1.0', 'end-1c'), encoding='utf-8')
            elif fmt == FMT_MD:
                out_path.write_text(build_md(self._data, doc_title), encoding='utf-8')
            elif fmt == FMT_DOCX:
                out_path.write_bytes(build_docx(self._data, doc_title))
            messagebox.showinfo('저장 완료', f'저장 완료:\n{out_path}')
            self._set_status(f'저장 완료 — {outname}')
        except Exception as e:
            messagebox.showerror('저장 오류', str(e))

    def _copy(self):
        self.clipboard_clear()
        fmt = self.v_fmt.get()
        if self._data and fmt != FMT_TXT:
            # DOCX/MD 서식 미리보기 상태 → TXT로 변환해서 복사
            content = build_txt(self._data, self.v_doc_title.get())
        else:
            content = self.txt.get('1.0', 'end-1c')
        self.clipboard_append(content)
        self._set_status('클립보드에 복사되었습니다. (TXT 형식)')

    def _set_status(self, msg: str):
        self.after(0, lambda: self.lbl_st.config(text=msg))


# ──────────────────────────── 진입점 ─────────────────────────────────────────

if __name__ == '__main__':
    app = App()
    app.mainloop()
